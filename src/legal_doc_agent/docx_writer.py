"""Markdown-ish to DOCX writer for generated legal packages."""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


@dataclass(frozen=True)
class DocumentSection:
    """A generated section to place in the final DOCX."""

    title: str
    markdown: str


def write_docx(
    *,
    title: str,
    subtitle: str,
    sections: list[DocumentSection],
    output_path: Path,
) -> None:
    """Write generated sections to a polished Word document."""

    document = Document()
    _configure_document(document)
    _add_cover(document, title, subtitle)

    for index, section in enumerate(sections):
        if index == 0:
            document.add_section(WD_SECTION.NEW_PAGE)
        else:
            document.add_page_break()
        document.add_heading(section.title, level=1)
        _add_markdown(document, section.markdown)

    document.save(output_path)


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in [
        ("Title", 22, RGBColor(31, 78, 121)),
        ("Subtitle", 12, RGBColor(89, 89, 89)),
        ("Heading 1", 15, RGBColor(31, 78, 121)),
        ("Heading 2", 12.5, RGBColor(55, 96, 146)),
        ("Heading 3", 11.5, RGBColor(64, 64, 64)),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def _add_cover(document: Document, title: str, subtitle: str) -> None:
    title_paragraph = document.add_paragraph(style="Title")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.add_run(title)

    subtitle_paragraph = document.add_paragraph(style="Subtitle")
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_paragraph.add_run(subtitle)

    notice = document.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = notice.add_run(
        "Drafting support output only. Review with qualified counsel before use."
    )
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(112, 48, 160)


def _add_markdown(document: Document, markdown: str) -> None:
    pending_table: list[list[str]] = []

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if _is_table_line(line):
            pending_table.append(_split_table_row(line))
            continue

        if pending_table:
            _flush_table(document, pending_table)
            pending_table = []

        stripped = line.strip()
        if not stripped:
            continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            level = min(len(heading.group(1)) + 1, 4)
            document.add_heading(_clean_inline_markdown(heading.group(2)).strip(), level=level)
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet:
            document.add_paragraph(
                _clean_inline_markdown(bullet.group(1)).strip(),
                style="List Bullet",
            )
            continue

        numbered = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if numbered:
            document.add_paragraph(
                _clean_inline_markdown(numbered.group(1)).strip(),
                style="List Number",
            )
            continue

        paragraph = document.add_paragraph()
        _add_runs_with_basic_markdown(paragraph, stripped)

    if pending_table:
        _flush_table(document, pending_table)


def _is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def _split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _flush_table(document: Document, rows: list[list[str]]) -> None:
    data_rows = [
        row
        for row in rows
        if not all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in row)
    ]
    if not data_rows:
        return

    column_count = max(len(row) for row in data_rows)
    table = document.add_table(rows=0, cols=column_count)
    table.style = "Table Grid"
    table.autofit = True

    for row_index, row_values in enumerate(data_rows):
        row = table.add_row()
        for cell_index in range(column_count):
            value = row_values[cell_index] if cell_index < len(row_values) else ""
            paragraph = row.cells[cell_index].paragraphs[0]
            _add_runs_with_basic_markdown(paragraph, value)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
                _set_cell_shading(row.cells[cell_index], "D9EAF7")


def _set_cell_shading(cell, fill: str) -> None:  # type: ignore[no-untyped-def]
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _add_runs_with_basic_markdown(paragraph, text: str) -> None:  # type: ignore[no-untyped-def]
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(_clean_inline_markdown(part))


def _clean_inline_markdown(text: str) -> str:
    return (
        text.replace("**", "")
        .replace("__", "")
        .replace("`", "")
        .replace("\\", "")
    )
