from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from legal_doc_agent.docx_writer import DocumentSection, write_docx


class DocxWriterTests(unittest.TestCase):
    def test_write_docx_creates_readable_document(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "sample.docx"

            write_docx(
                title="Package",
                subtitle="Subtitle",
                sections=[
                    DocumentSection(
                        title="PART A",
                        markdown="# Heading\n\n- Bullet\n\n1. Numbered\n\nPlain **bold** text.",
                    )
                ],
                output_path=output_path,
            )

            self.assertTrue(output_path.exists())
            document = Document(output_path)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("Package", text)
            self.assertIn("PART A", text)
            self.assertIn("Bullet", text)
            self.assertIn("Plain bold text.", text)

    def test_write_docx_removes_markdown_code_fence(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "sample.docx"

            write_docx(
                title="Package",
                subtitle="Subtitle",
                sections=[
                    DocumentSection(
                        title="PART A",
                        markdown="```markdown\n# Planner Summary\n\n**Matter Type:** Founder package\n```",
                    )
                ],
                output_path=output_path,
            )

            document = Document(output_path)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("Planner Summary", text)
            self.assertIn("Matter Type: Founder package", text)
            self.assertNotIn("```", text)


if __name__ == "__main__":
    unittest.main()
