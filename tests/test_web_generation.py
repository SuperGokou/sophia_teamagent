from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from legal_doc_agent.web_generation import generate_web_legal_package


class TruncatedClient:
    def complete(
        self,
        messages: list[dict[str, str]],
        *,
        role: str | None = None,
    ) -> str:
        return (
            "# Planner Summary\n\n"
            "Matter type: Delaware post-formation package.\n\n"
            "# Draft Package\n\n"
            "1. Contractor Agreement: Purpose: Define contractor services. Prep: Scope of"
        )


class TimeoutThenSuccessClient:
    def __init__(self) -> None:
        self.calls: list[list[dict[str, str]]] = []

    def complete(
        self,
        messages: list[dict[str, str]],
        *,
        role: str | None = None,
    ) -> str:
        self.calls.append(messages)
        if len(self.calls) == 1:
            raise TimeoutError("The read operation timed out")
        return (
            "# Planner Summary\n\n"
            "Matter type: Delaware post-formation package.\n\n"
            "# Draft Package\n\n"
            "Founder agreement and board consent package with placeholders.\n\n"
            "# Reviewer Quality Gate\n\n"
            "Counsel must review all placeholders and filing dates.\n\n"
            "END OF PACKAGE"
        )


class CapturingClient:
    def __init__(self) -> None:
        self.messages: list[dict[str, str]] = []

    def complete(
        self,
        messages: list[dict[str, str]],
        *,
        role: str | None = None,
    ) -> str:
        self.messages = messages
        return (
            "# Planner Summary\n\n"
            "Matter type: Delaware founder package.\n\n"
            "# Draft Package\n\n"
            "Use counsel-reviewed placeholders.\n\n"
            "# Reviewer Quality Gate\n\n"
            "Check retrieved citations and version dates.\n\n"
            "END OF PACKAGE"
        )


class AlwaysTimeoutClient:
    def __init__(self) -> None:
        self.calls: list[list[dict[str, str]]] = []

    def complete(
        self,
        messages: list[dict[str, str]],
        *,
        role: str | None = None,
    ) -> str:
        self.calls.append(messages)
        raise TimeoutError("The read operation timed out")


LONG_POST_FORMATION_PROMPT = (
    "You are a top-tier Silicon Valley startup attorney and Delaware corporate counsel. "
    "Generate a complete institutional-grade post-formation legal documentation package "
    "for a Delaware C-Corporation startup. Do not summarize excessively. Do not omit key clauses. "
    "Continue sequentially until all required documents are complete.\n\n"
    "# COMPANY PROFILE\n"
    "Company Type: AI software company + AI agent staffing platform\n"
    "Business Activities: AI agents, AI-powered staffing systems, SaaS products, enterprise productivity software\n"
    "Jurisdiction: Delaware C-Corporation\n"
    "Founders: 2 founders\n"
    "Founder Ownership: 50/50\n"
    "Authorized Shares: 10,000,000 common shares\n"
    "Founder Allocation: 5,000,000 shares each\n"
    "Founder Vesting: 4-year vesting, 1-year cliff, monthly vesting thereafter\n\n"
    "# STEP 1 REQUIRED VS OPTIONAL DOCUMENT CHECKLIST\n"
    "Required documents include corporate bylaws, initial board consent, founder stock purchase agreements, "
    "stock ledger, cap table, IP assignment, CIIAA, 83(b) election instructions, banking authorization, "
    "officer appointment, founder vesting and repurchase rights.\n\n"
    "# STEP 3 LAW-FIRM-GRADE DOCUMENT TEMPLATES\n"
    "For every required document generate a complete, detailed, professional template drafted in the style "
    "of Cooley, Wilson Sonsini, Gunderson, Orrick, YC and Clerky startup standards.\n\n"
    + ("Repeat complete professional template requirement. " * 140)
)


class WebGenerationTests(unittest.TestCase):
    def test_truncated_generation_gets_completion_safeguard(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            result = generate_web_legal_package(
                client=TruncatedClient(),
                brief="Generate a Delaware founder legal package.",
                output_path=root / "package.docx",
                artifact_dir=root / "artifacts",
            )

            artifact_text = (result.artifact_dir / "web_drafter_package.md").read_text(
                encoding="utf-8"
            )
            self.assertIn("# Reviewer Quality Gate", artifact_text)
            self.assertIn("Completion Safeguard", artifact_text)
            self.assertNotIn("Scope of\n", artifact_text)

            doc_text = "\n".join(
                paragraph.text for paragraph in Document(result.output_path).paragraphs
            )
            self.assertIn("Reviewer Quality Gate", doc_text)
            self.assertIn("Completion Safeguard", doc_text)

    def test_provider_timeout_retries_with_short_package_prompt(self) -> None:
        client = TimeoutThenSuccessClient()
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            result = generate_web_legal_package(
                client=client,
                brief="Generate a Delaware founder legal package.",
                output_path=root / "package.docx",
                artifact_dir=root / "artifacts",
            )

            artifact_text = (result.artifact_dir / "web_drafter_package.md").read_text(
                encoding="utf-8"
            )
            retry_prompt = client.calls[1][1]["content"]
            self.assertEqual(len(client.calls), 2)
            self.assertIn("Target 700-1,100 words", retry_prompt)
            self.assertIn("Founder agreement and board consent package", artifact_text)
            self.assertTrue(result.output_path.exists())

    def test_web_generation_includes_knowledge_context_in_prompt(self) -> None:
        client = CapturingClient()
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            result = generate_web_legal_package(
                client=client,
                brief="Generate a Delaware founder legal package.",
                output_path=root / "package.docx",
                artifact_dir=root / "artifacts",
                knowledge_context=(
                    "[1] 8 Del. C. § 141(f) - Board action without meeting\n"
                    "version_date: 2026-06-06\n"
                    "excerpt: Board action may be taken by consent."
                ),
            )
            artifact_text = (result.artifact_dir / "web_drafter_package.md").read_text(
                encoding="utf-8"
            )

        prompt = client.messages[1]["content"]
        self.assertIn("SUPPLEMENTAL LEGAL KNOWLEDGE BASE CONTEXT", prompt)
        self.assertIn("8 Del. C. § 141(f)", prompt)
        self.assertIn("version_date: 2026-06-06", prompt)
        self.assertIn("Do not invent citations", prompt)
        self.assertIn("# Retrieved Authority Context", artifact_text)
        self.assertIn("8 Del. C. § 141(f)", artifact_text)
        self.assertIn("version_date: 2026-06-06", artifact_text)

    def test_long_prompt_is_compacted_before_provider_call(self) -> None:
        client = CapturingClient()
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            generate_web_legal_package(
                client=client,
                brief=LONG_POST_FORMATION_PROMPT,
                output_path=root / "package.docx",
                artifact_dir=root / "artifacts",
            )

        prompt = client.messages[1]["content"]
        self.assertLess(len(prompt), 5000)
        self.assertIn("ONLINE-SAFE REQUEST DIGEST", prompt)
        self.assertIn("Delaware C-Corporation", prompt)
        self.assertIn("83(b) election instructions", prompt)
        self.assertNotIn("Repeat complete professional template requirement", prompt)

    def test_double_provider_timeout_returns_recovery_docx(self) -> None:
        client = AlwaysTimeoutClient()
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            result = generate_web_legal_package(
                client=client,
                brief=LONG_POST_FORMATION_PROMPT,
                output_path=root / "package.docx",
                artifact_dir=root / "artifacts",
            )
            artifact_text = (result.artifact_dir / "web_drafter_package.md").read_text(
                encoding="utf-8"
            )
            doc_text = "\n".join(
                paragraph.text for paragraph in Document(result.output_path).paragraphs
            )
            self.assertTrue(result.output_path.exists())

        self.assertEqual(len(client.calls), 2)
        self.assertIn("AI Provider Timeout Recovery Package", artifact_text)
        self.assertIn("Required Document Checklist", artifact_text)
        self.assertIn("END OF PACKAGE", artifact_text)
        self.assertEqual(result.generation_mode, "timeout_recovery")
        self.assertNotIn("NVIDIA", artifact_text)
        self.assertNotIn("NVIDIA", doc_text)

    def test_generated_artifacts_do_not_include_provider_branding(self) -> None:
        client = CapturingClient()
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            result = generate_web_legal_package(
                client=client,
                brief="Generate a Delaware founder legal package.",
                output_path=root / "package.docx",
                artifact_dir=root / "artifacts",
            )
            artifact_text = (result.artifact_dir / "web_drafter_package.md").read_text(
                encoding="utf-8"
            )
            doc_text = "\n".join(
                paragraph.text for paragraph in Document(result.output_path).paragraphs
            )

        self.assertNotIn("NVIDIA", artifact_text)
        self.assertNotIn("NVIDIA", doc_text)


if __name__ == "__main__":
    unittest.main()
